from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
import os

# === CONFIGURATION ===
pdf_path = "A-06.pdf"     # your input PDF file
output_docx = "A-06.docx"
poppler_path = r"C:\Program Files\poppler-25.07.0\Library\bin"

# e.g., poppler_path = r"C:\Program Files\poppler-24.02.0\Library\bin"

# === STEP 1: Convert PDF pages to images ===
print("📄 Converting PDF pages to images...")
pages = convert_from_path(pdf_path, dpi=300, poppler_path=poppler_path)

# === STEP 2: Create new Word document ===
doc = Document()

# === STEP 3: Process each page ===
for i, page in enumerate(pages, start=1):
    print(f"🔍 Processing page {i}...")
    # Extract text using OCR
    text = pytesseract.image_to_string(page)
    
    # Add to Word document
    doc.add_paragraph(f"--- Page {i} ---")
    doc.add_paragraph(text)
    doc.add_page_break()

# === STEP 4: Save DOCX file ===
doc.save(output_docx)
print(f"✅ Conversion complete! Saved as '{output_docx}'")




# writer: sajid_islam